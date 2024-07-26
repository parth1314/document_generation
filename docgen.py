import mysql.connector
import os
from dotenv import load_dotenv
from jinja2 import Template
from docx import Document

load_dotenv()

#RDS DB Credentials
host_name = os.getenv('HOST_NAME')
user_name = os.getenv('USER')
password = os.getenv('PASSWORD')
db_name = os.getenv('DATABASE')

conn = mysql.connector.connect(
    host=host_name,
    user=user_name,
    password=password,
    database=db_name
)
cursor = conn.cursor(dictionary=True)


cursor.execute("SELECT * FROM ra_pci_dss_req")
pci_dss_req = cursor.fetchall()

cursor.execute("SELECT * FROM ra_appendix_c_data_kbox")
appendix_c = cursor.fetchall()

cursor.execute("SELECT * FROM ra_appendix_e_data_kbox")
appendix_e = cursor.fetchall()

conn.close()


num_rows = len(pci_dss_req)
if num_rows != len(appendix_c) or num_rows != len(appendix_e):
    print("Error: The number of rows in the tables are not equal.")
    exit()


template_path = r'C:\Users\pagrawal\Downloads\document_generation\Modified_R1_Template - Copy.docx'

def add_checkbox(doc, placeholder, is_checked):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, '☒' if is_checked else '☐')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, '☒ ' if is_checked else '☐ ')


def replace_placeholder(doc, placeholder, value):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))


for i in range(num_rows):
    doc = Document(template_path)
    for key, value in pci_dss_req[i].items():
        replace_placeholder(doc, f'{{{{ pci_dss_req.{key} }}}}', value)

    add_checkbox(doc, f'{{{{ pci_dss_req.checkbox_is_in_place }}}}', pci_dss_req[i].get('is_in_place') == 'YES')
    add_checkbox(doc, f'{{{{ pci_dss_req.checkbox_is_not_applicable }}}}', pci_dss_req[i].get('is_not_applicable') == 'YES')
    add_checkbox(doc, f'{{{{ pci_dss_req.checkbox_is_not_tested }}}}', pci_dss_req[i].get('is_not_tested') == 'YES')
    add_checkbox(doc, f'{{{{ pci_dss_req.checkbox_is_not_in_place }}}}', pci_dss_req[i].get('is_not_in_place') == 'YES')

    
    if pci_dss_req[i].get('indicate_whether_compensating_control_is_used') == 'YES':
        for key, value in appendix_c[i].items():
            replace_placeholder(doc, f'{{{{ appendix_c.{key} }}}}', value)
        add_checkbox(doc, f'{{{{ appendix_c.checkbox_yes }}}}', pci_dss_req[i].get('indicate_whether_compensating_control_is_used') == 'YES')
        add_checkbox(doc, f'{{{{ appendix_c.checkbox_no }}}}', pci_dss_req[i].get('indicate_whether_compensating_control_is_used') == 'NO')
    else:
        print(f"Skipping Appendix C data for row {i+1}...")
        for key in appendix_c[i].keys():
            replace_placeholder(doc, f'{{{{ appendix_c.{key} }}}}', '')
        add_checkbox(doc, f'{{{{ appendix_c.checkbox_yes }}}}', False)
        add_checkbox(doc, f'{{{{ appendix_c.checkbox_no }}}}', True)

    
    if pci_dss_req[i].get('indicate_whether_customised_approach_is_used') == 'YES':
        for key, value in appendix_e[i].items():
            replace_placeholder(doc, f'{{{{ appendix_e.{key} }}}}', value)
        add_checkbox(doc, f'{{{{ appendix_e.checkbox_yes }}}}', pci_dss_req[i].get('indicate_whether_customised_approach_is_used') == 'YES')
        add_checkbox(doc, f'{{{{ appendix_e.checkbox_no }}}}', pci_dss_req[i].get('indicate_whether_customised_approach_is_used') == 'NO')
    else:
        for key in appendix_e[i].keys():
            replace_placeholder(doc, f'{{{{ appendix_e.{key} }}}}', '')
        add_checkbox(doc, f'{{{{ appendix_e.checkbox_yes }}}}', False)
        add_checkbox(doc, f'{{{{ appendix_e.checkbox_no }}}}', True)

    
    output_path = f'C:\\Users\\pagrawal\\Downloads\\document_generation\\output_filled_template_{i+1}.docx'
    doc.save(output_path)

    print(f"Generated document saved to {output_path}")

print("Document generation completed.")
